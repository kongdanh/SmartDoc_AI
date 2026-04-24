"""
SmartDoc AI — Document Preprocessor.

Converts PDF / DOCX files to plain text (.txt).
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


def convert_pdf_to_txt(pdf_path: Path, txt_path: Optional[Path] = None) -> Optional[Path]:
    """
    Extract text from a PDF and write it to a .txt file.

    Args:
        pdf_path: Path to the source PDF.
        txt_path: Where to save the output. Defaults to same dir, .txt extension.

    Returns:
        Path to the generated .txt, or None on failure.
    """
    try:
        import fitz  # PyMuPDF

        if txt_path is None:
            txt_path = pdf_path.with_suffix(".txt")

        doc = fitz.open(str(pdf_path))
        text_parts = []
        for page in doc:
            page_text = page.get_text("text") or ""
            page_text = page_text.strip()
            if page_text:
                text_parts.append(page_text)
        doc.close()

        full_text = "\n\n".join(text_parts)
        if not full_text.strip():
            logger.warning("PDF has no extractable text: %s", pdf_path.name)
            return None

        txt_path.write_text(full_text, encoding="utf-8")
        logger.info("Converted PDF → TXT: %s (%d chars)", txt_path.name, len(full_text))
        return txt_path

    except Exception as e:
        logger.error("Failed to convert PDF %s: %s", pdf_path.name, e)
        return None


def convert_docx_to_txt(docx_path: Path, txt_path: Optional[Path] = None) -> Optional[Path]:
    """
    Extract text from a DOCX and write it to a .txt file.

    Args:
        docx_path: Path to the source DOCX.
        txt_path: Where to save the output. Defaults to same dir, .txt extension.

    Returns:
        Path to the generated .txt, or None on failure.
    """
    try:
        from docx import Document

        if txt_path is None:
            txt_path = docx_path.with_suffix(".txt")

        doc = Document(str(docx_path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        full_text = "\n\n".join(paragraphs)
        if not full_text.strip():
            logger.warning("DOCX has no extractable text: %s", docx_path.name)
            return None

        txt_path.write_text(full_text, encoding="utf-8")
        logger.info("Converted DOCX → TXT: %s (%d chars)", txt_path.name, len(full_text))
        return txt_path

    except Exception as e:
        logger.error("Failed to convert DOCX %s: %s", docx_path.name, e)
        return None


def read_txt_file(txt_path: Path) -> Optional[str]:
    """Read a plain text file and return its content."""
    try:
        return txt_path.read_text(encoding="utf-8")
    except Exception as e:
        logger.error("Failed to read TXT %s: %s", txt_path.name, e)
        return None
