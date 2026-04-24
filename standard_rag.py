"""
SmartDoc AI — Standard RAG Engine.

Self-contained RAG pipeline (ChromaDB + LLM) that can run independently.
Adapted from the RAG/ directory code but uses OpenRouter LLM
(same key as GraphRAG) for universal compatibility.

Usage from main.py:
    from standard_rag import build_standard_rag_index, query_standard_rag
"""

from __future__ import annotations

import hashlib
import json
import logging
import os
import re
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from dotenv import load_dotenv
from langchain_community.embeddings import OpenAIEmbeddings

load_dotenv()

logger = logging.getLogger(__name__)

# ── Configuration ───────────────────────────────────────────────

_ROOT = Path(__file__).resolve().parent
_DB_DIR = _ROOT / "db" / "standard_rag_chroma"
_DATA_DIR = _ROOT / "data"

# LLM settings (OpenRouter — same as GraphRAG)
_LLM_BASE_URL = os.getenv("LLM_BASE_URL", "https://openrouter.ai/api/v1")
_LLM_API_KEY = os.getenv("LLM_API_KEY", "")
_LLM_MODEL = os.getenv("LLM_MODEL", "meta-llama/llama-3.3-70b-instruct:free")

# Embedding (HuggingFace local)
_EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "sentence-transformers/all-MiniLM-L6-v2")


# ═══════════════════════════════════════════════════════════════
# EMBEDDING (HuggingFace — local, free)
# ═══════════════════════════════════════════════════════════════

_embedder_instance = None


from langchain_openai import OpenAIEmbeddings

class HuggingFaceEmbedder:
    """Embedding using OpenRouter/OpenAI."""

    def __init__(self, model_name=None):
        if model_name is None:
            model_name = _EMBEDDING_MODEL
        api_key = os.getenv("LLM_API_KEY")
        base_url = os.getenv("LLM_BASE_URL", "https://openrouter.ai/api/v1")
        
        self.embeddings = OpenAIEmbeddings(
            model=model_name,
            api_key=api_key,
            base_url=base_url 
        )

    def embed_text(self, text: str) -> List[float]:
        text = (text or "").strip()
        if not text:
            return []
        return self.embeddings.embed_query(text)

    def embed_texts(self, texts: List[str]) -> List[List[float]]:
        texts = [t.strip() for t in texts if t.strip()]
        if not texts:
            return []
        return self.embeddings.embed_documents(texts)


def _get_embedder() -> HuggingFaceEmbedder:
    global _embedder_instance
    if _embedder_instance is None:
        _embedder_instance = HuggingFaceEmbedder()
    return _embedder_instance


# ═══════════════════════════════════════════════════════════════
# DOCUMENT LOADING
# ═══════════════════════════════════════════════════════════════


def _load_document(file_path: Path) -> List[Dict[str, Any]]:
    """
    Load a document (PDF, DOCX, TXT) and return list of page dicts.
    Each dict: { "page": int, "text": str, "metadata": {...} }
    """
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return _load_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        return _load_docx(file_path)
    elif suffix == ".txt":
        return _load_txt(file_path)
    else:
        logger.warning("Unsupported file type: %s", suffix)
        return []


def _load_pdf(file_path: Path) -> List[Dict[str, Any]]:
    """Extract text from PDF using PyMuPDF."""
    try:
        import fitz
        doc = fitz.open(str(file_path))
        pages = []
        for i, page in enumerate(doc):
            text = page.get_text("text") or ""
            text = text.replace("\r", "\n").strip()
            if text:
                pages.append({
                    "page": i + 1,
                    "text": text,
                    "metadata": {
                        "source": file_path.name,
                        "page": i + 1,
                        "document_type": "pdf",
                    },
                })
        doc.close()
        return pages
    except Exception as e:
        logger.error("Failed to load PDF %s: %s", file_path.name, e)
        return []


def _load_docx(file_path: Path) -> List[Dict[str, Any]]:
    """Extract text from DOCX."""
    try:
        from docx import Document
        doc = Document(str(file_path))

        paragraphs = []
        for p in doc.paragraphs:
            if p.text.strip():
                paragraphs.append(p.text.strip())

        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        full_text = "\n\n".join(paragraphs)
        if not full_text.strip():
            return []

        return [{
            "page": 1,
            "text": full_text,
            "metadata": {
                "source": file_path.name,
                "page": 1,
                "document_type": "docx",
            },
        }]
    except Exception as e:
        logger.error("Failed to load DOCX %s: %s", file_path.name, e)
        return []


def _load_txt(file_path: Path) -> List[Dict[str, Any]]:
    """Load plain text file."""
    try:
        text = file_path.read_text(encoding="utf-8")
        if not text.strip():
            return []
        return [{
            "page": 1,
            "text": text.strip(),
            "metadata": {
                "source": file_path.name,
                "page": 1,
                "document_type": "txt",
            },
        }]
    except Exception as e:
        logger.error("Failed to load TXT %s: %s", file_path.name, e)
        return []


# ═══════════════════════════════════════════════════════════════
# CHUNKING
# ═══════════════════════════════════════════════════════════════


def _chunk_text(
    text: str,
    chunk_size: int = 500,
    overlap: int = 50,
) -> List[str]:
    """Split text into overlapping chunks by sentences."""
    if not text.strip():
        return []

    # Split by sentences
    sentences = re.split(r'(?<=[.!?。])\s+|\n\n+', text)
    sentences = [s.strip() for s in sentences if s.strip()]

    if not sentences:
        return [text.strip()] if text.strip() else []

    chunks = []
    current_chunk = []
    current_length = 0

    for sentence in sentences:
        sentence_words = len(sentence.split())

        if current_length + sentence_words > chunk_size and current_chunk:
            chunks.append(" ".join(current_chunk))

            # Keep overlap
            overlap_words = 0
            overlap_start = len(current_chunk)
            for j in range(len(current_chunk) - 1, -1, -1):
                overlap_words += len(current_chunk[j].split())
                if overlap_words >= overlap:
                    overlap_start = j
                    break
            current_chunk = current_chunk[overlap_start:]
            current_length = sum(len(s.split()) for s in current_chunk)

        current_chunk.append(sentence)
        current_length += sentence_words

    if current_chunk:
        chunks.append(" ".join(current_chunk))

    return chunks


def _build_chunks_from_pages(
    pages: List[Dict[str, Any]],
    file_name: str,
) -> List[Dict[str, Any]]:
    """Build indexed chunks from loaded pages."""
    all_chunks = []
    chunk_idx = 0

    for page_data in pages:
        text = page_data.get("text", "")
        page_num = page_data.get("page", 0)
        source_meta = page_data.get("metadata", {})

        text_chunks = _chunk_text(text)

        for chunk_text in text_chunks:
            chunk_id = f"std_rag_{hashlib.md5(f'{file_name}:{page_num}:{chunk_idx}'.encode()).hexdigest()[:12]}"

            all_chunks.append({
                "id": chunk_id,
                "content": chunk_text,
                "metadata": {
                    "source": file_name,
                    "page": page_num,
                    "chunk_index": chunk_idx,
                    **source_meta,
                },
            })
            chunk_idx += 1

    return all_chunks


# ═══════════════════════════════════════════════════════════════
# VECTOR STORE (ChromaDB)
# ═══════════════════════════════════════════════════════════════

_chroma_clients: Dict[str, Any] = {}


def _get_chroma_collection(domain: str = "default"):
    """Get or create a ChromaDB collection for a domain."""
    import chromadb

    if domain not in _chroma_clients:
        db_path = str(_DB_DIR / domain)
        Path(db_path).mkdir(parents=True, exist_ok=True)
        client = chromadb.PersistentClient(path=db_path)
        collection = client.get_or_create_collection(
            name=f"rag_{domain}",
            metadata={"hnsw:space": "cosine"},
        )
        _chroma_clients[domain] = {"client": client, "collection": collection}

    return _chroma_clients[domain]["collection"]


def _reset_collection(domain: str = "default"):
    """Reset a domain's collection."""
    import chromadb

    db_path = str(_DB_DIR / domain)
    Path(db_path).mkdir(parents=True, exist_ok=True)
    client = chromadb.PersistentClient(path=db_path)

    collection_name = f"rag_{domain}"
    try:
        client.delete_collection(name=collection_name)
    except Exception:
        pass

    collection = client.get_or_create_collection(
        name=collection_name,
        metadata={"hnsw:space": "cosine"},
    )
    _chroma_clients[domain] = {"client": client, "collection": collection}
    return collection


# ═══════════════════════════════════════════════════════════════
# INDEXING (Build vector database)
# ═══════════════════════════════════════════════════════════════


def build_standard_rag_index(domain: str = "default") -> Dict[str, Any]:
    """
    Build/rebuild the Standard RAG index for a domain.

    Reads all files in data/<domain>/ directory, chunks them,
    embeds with HuggingFace, stores in ChromaDB.
    """
    domain_data_dir = _DATA_DIR / domain
    if not domain_data_dir.exists():
        return {
            "status": "error",
            "message": f"Domain directory not found: {domain_data_dir}",
            "files": 0,
            "chunks": 0,
        }

    # Find all supported files
    file_extensions = {".pdf", ".docx", ".doc", ".txt"}
    files = [
        f for f in domain_data_dir.rglob("*")
        if f.is_file() and f.suffix.lower() in file_extensions
    ]

    if not files:
        return {
            "status": "warning",
            "message": f"No supported files found in {domain}",
            "files": 0,
            "chunks": 0,
        }

    logger.info("Building Standard RAG index for domain '%s' (%d files)", domain, len(files))

    # Reset collection
    collection = _reset_collection(domain)
    embedder = _get_embedder()

    total_chunks = 0

    for file_path in files:
        try:
            # Load document
            pages = _load_document(file_path)
            if not pages:
                logger.warning("No content from %s, skipping", file_path.name)
                continue

            # Build chunks
            chunks = _build_chunks_from_pages(pages, file_path.name)
            if not chunks:
                continue

            # Embed and store in batches
            batch_size = 50
            for i in range(0, len(chunks), batch_size):
                batch = chunks[i:i + batch_size]

                texts = [c["content"] for c in batch]
                embeddings = embedder.embed_texts(texts)

                if not embeddings:
                    continue

                ids = [c["id"] for c in batch]
                documents = texts
                metadatas = []
                for c in batch:
                    # Sanitize metadata for ChromaDB (only str, int, float, bool)
                    meta = {}
                    for k, v in c["metadata"].items():
                        if isinstance(v, (str, int, float, bool)):
                            meta[k] = v
                        else:
                            meta[k] = str(v)
                    metadatas.append(meta)

                collection.add(
                    ids=ids,
                    documents=documents,
                    embeddings=embeddings,
                    metadatas=metadatas,
                )

                total_chunks += len(batch)

            logger.info("Indexed %s: %d chunks", file_path.name, len(chunks))

        except Exception as e:
            logger.error("Error indexing %s: %s", file_path.name, e)

    return {
        "status": "success",
        "message": f"Standard RAG index built for '{domain}'",
        "files": len(files),
        "chunks": total_chunks,
    }


def build_faiss_index(txt_file: Path) -> bool:
    """
    Legacy compatibility: Build index from a single text file.
    Extracts the domain from the file's parent directory.
    """
    try:
        # Determine domain from file path
        domain = txt_file.parent.name
        if domain in ("pdf", "docx", "txt", "raw", "document"):
            # Navigate up to find the actual domain
            for parent in txt_file.parents:
                if parent.parent == _DATA_DIR:
                    domain = parent.name
                    break

        # Just trigger a full domain rebuild
        result = build_standard_rag_index(domain)
        return result.get("status") == "success"
    except Exception as e:
        logger.error("build_faiss_index error: %s", e)
        return False


# ═══════════════════════════════════════════════════════════════
# QUERYING
# ═══════════════════════════════════════════════════════════════


def query_standard_rag(
    query: str,
    domain: str = "default",
    top_k: int = 8,
) -> str:
    """
    Query the Standard RAG pipeline.

    1. Embed the query
    2. Retrieve relevant chunks from ChromaDB
    3. Build context
    4. Generate answer via LLM (OpenRouter)

    Returns:
        Answer string.
    """
    query = (query or "").strip()
    if not query:
        return "Vui lòng nhập câu hỏi."

    # 1. Get collection
    try:
        collection = _get_chroma_collection(domain)
        count = collection.count()
        if count == 0:
            return f"Chưa có dữ liệu trong domain '{domain}'. Hãy upload và index tài liệu trước."
    except Exception as e:
        return f"Lỗi kết nối ChromaDB: {str(e)}"

    # 2. Embed query
    embedder = _get_embedder()
    query_embedding = embedder.embed_text(query)
    if not query_embedding:
        return "Lỗi: Không thể tạo embedding cho câu hỏi."

    # 3. Retrieve relevant documents
    try:
        results = collection.query(
            query_embeddings=[query_embedding],
            n_results=min(top_k, count),
        )
    except Exception as e:
        return f"Lỗi truy vấn ChromaDB: {str(e)}"

    documents = results.get("documents", [[]])[0]
    metadatas = results.get("metadatas", [[]])[0]

    if not documents:
        return "Không tìm thấy tài liệu liên quan trong cơ sở dữ liệu."

    # 4. Build context
    context = _build_context(documents, metadatas)

    # 5. Generate answer
    answer = _generate_answer(query, context)
    return answer


def _build_context(documents: List[str], metadatas: List[Dict]) -> str:
    """Build context string from retrieved documents."""
    parts = []
    for i, (doc, meta) in enumerate(zip(documents, metadatas), start=1):
        source = meta.get("source", "unknown")
        page = meta.get("page", "N/A")
        parts.append(f"[{i}] {source}, trang {page}\n{doc}")
    return "\n\n".join(parts)


def _generate_answer(query: str, context: str) -> str:
    """Generate answer using LLM via OpenRouter API."""
    import requests

    prompt = f"""Bạn là trợ lý AI chuyên nghiệp, hỗ trợ tra cứu thông tin từ tài liệu.

NGUYÊN TẮC:
1. Phân tích câu hỏi và tổng hợp thông tin từ các đoạn tài liệu được cung cấp.
2. Trả lời bằng tiếng Việt, rõ ràng, chuyên nghiệp.
3. Nếu tài liệu không chứa thông tin liên quan, nói rõ "Không tìm thấy thông tin trong tài liệu".
4. Trích dẫn nguồn bằng [1], [2]... khi sử dụng thông tin từ tài liệu.

[CÂU HỎI]
{query}

[NỘI DUNG TÀI LIỆU]
{context}

Hãy trả lời:"""

    try:
        response = requests.post(
            f"{_LLM_BASE_URL}/chat/completions",
            headers={
                "Authorization": f"Bearer {_LLM_API_KEY}",
                "Content-Type": "application/json",
            },
            json={
                "model": _LLM_MODEL,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.3,
                "max_tokens": 2048,
            },
            timeout=120,
        )

        if response.status_code == 200:
            data = response.json()
            answer = data["choices"][0]["message"]["content"]
            return answer.strip()
        else:
            return f"Lỗi LLM API ({response.status_code}): {response.text[:200]}"

    except Exception as e:
        return f"Lỗi kết nối LLM: {str(e)}"
